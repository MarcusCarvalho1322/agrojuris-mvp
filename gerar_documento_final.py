# -*- coding: utf-8 -*-
"""
Gerador do Dossiê Agrodefesa Legal - Versão Transparente
Gera documento Word formatado profissionalmente
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def set_cell_shading(cell, color):
    """Define cor de fundo de uma célula"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_heading(doc, text, level=1):
    """Adiciona heading formatado"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0, 100, 0)  # Verde escuro
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 80, 0)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(50, 50, 50)
    
    return p

def add_warning_box(doc, text):
    """Adiciona caixa de aviso destacada"""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF3CD')  # Amarelo claro
    
    p = cell.paragraphs[0]
    run = p.add_run("⚠️ " + text)
    run.bold = True
    run.font.size = Pt(11)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()

def add_estimate_tag(paragraph):
    """Adiciona tag [ESTIMATIVA] formatada"""
    run = paragraph.add_run(" [ESTIMATIVA]")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 100, 0)  # Laranja
    run.font.size = Pt(9)

def create_document():
    doc = Document()
    
    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # =========================================================================
    # CAPA
    # =========================================================================
    for _ in range(3):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AGRODEFESA LEGAL")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0, 100, 0)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Plataforma de Defesa Regulatória para o Agronegócio")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run("DOSSIÊ DE INVESTIMENTO")
    run.bold = True
    run.font.size = Pt(20)
    
    doc.add_paragraph()
    
    transparent = doc.add_paragraph()
    transparent.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = transparent.add_run("Versão Transparente para Investidores")
    run.italic = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    for _ in range(5):
        doc.add_paragraph()
    
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(f"Janeiro 2026")
    run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # =========================================================================
    # AVISO AO INVESTIDOR
    # =========================================================================
    add_formatted_heading(doc, "AVISO IMPORTANTE AO INVESTIDOR", 1)
    doc.add_paragraph()
    
    warning_table = doc.add_table(rows=1, cols=1)
    warning_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = warning_table.rows[0].cells[0]
    set_cell_shading(cell, 'FFF3CD')
    
    p = cell.paragraphs[0]
    p.add_run("\n")
    run = p.add_run("Este documento foi preparado com TRANSPARÊNCIA TOTAL.\n\n")
    run.bold = True
    
    p.add_run("Ele contém uma combinação de:\n\n")
    
    run = p.add_run("✓ DADOS VERIFICÁVEIS: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    p.add_run("Obtidos de fontes públicas oficiais\n\n")
    
    run = p.add_run("⚠ ESTIMATIVAS: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 150, 0)
    p.add_run("Projeções baseadas em metodologia documentada\n\n")
    
    run = p.add_run("? PREMISSAS: ")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 100, 0)
    p.add_run("Hipóteses que requerem validação de mercado\n\n")
    
    p.add_run("Todos os dados estimados estão claramente marcados e detalhados no Anexo.\n")
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # =========================================================================
    # RESUMO EXECUTIVO
    # =========================================================================
    add_formatted_heading(doc, "RESUMO EXECUTIVO", 1)
    doc.add_paragraph()
    
    # Problema
    add_formatted_heading(doc, "O PROBLEMA", 2)
    
    p = doc.add_paragraph()
    p.add_run("O agronegócio brasileiro enfrenta um volume significativo de autuações regulatórias nas esferas ambiental, trabalhista e sanitária. ")
    run = p.add_run("Este é um FATO VERIFICÁVEL.")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 128, 0)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Dados verificáveis:").bold = True
    
    bullets = [
        "IBAMA registra milhares de autos de infração anualmente (público via LAI)",
        "Código Florestal (Lei 12.651/2012) exige CAR e regularização ambiental",
        "Fiscalização trabalhista rural documentada pelo MTE/SIT",
        "MAPA realiza fiscalizações sanitárias com dados parcialmente públicos"
    ]
    
    for bullet in bullets:
        p = doc.add_paragraph(bullet, style='List Bullet')
    
    doc.add_paragraph()
    
    # Solução
    add_formatted_heading(doc, "A SOLUÇÃO", 2)
    
    p = doc.add_paragraph()
    p.add_run("Plataforma tecnológica para defesa regulatória agrícola:")
    
    solutions = [
        "Monitoramento proativo de riscos regulatórios",
        "Alertas preventivos de fiscalização",
        "Suporte técnico-jurídico para defesas administrativas",
        "Documentação e compliance automatizado"
    ]
    
    for sol in solutions:
        doc.add_paragraph(sol, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # =========================================================================
    # TAMANHO DE MERCADO
    # =========================================================================
    add_formatted_heading(doc, "TAMANHO DE MERCADO", 1)
    doc.add_paragraph()
    
    add_warning_box(doc, "Os valores abaixo são ESTIMATIVAS baseadas em metodologia documentada no Anexo")
    
    # Tabela TAM/SAM/SOM
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Cabeçalho
    header_cells = table.rows[0].cells
    for i, text in enumerate(["Métrica", "Valor", "Status"]):
        header_cells[i].text = text
        header_cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(header_cells[i], '2E7D32')
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    # TAM
    row = table.rows[1].cells
    row[0].text = "TAM (Mercado Total)"
    row[1].text = "R$ 47,2 bilhões"
    row[2].text = "[ESTIMATIVA]"
    set_cell_shading(row[2], 'FFF3CD')
    
    # SAM
    row = table.rows[2].cells
    row[0].text = "SAM (Mercado Disponível)"
    row[1].text = "R$ 14,16 bilhões"
    row[2].text = "[ESTIMATIVA]"
    set_cell_shading(row[2], 'FFF3CD')
    
    # SOM
    row = table.rows[3].cells
    row[0].text = "SOM (Ano 1)"
    row[1].text = "R$ 283 milhões"
    row[2].text = "[ESTIMATIVA]"
    set_cell_shading(row[2], 'FFF3CD')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("Ressalva: ")
    run.bold = True
    p.add_run("Estes números são projeções baseadas em extrapolação de dados amostrais do IBAMA e estimativas para SEMAs estaduais. Os valores reais podem variar significativamente. Consulte o Anexo para metodologia completa.")
    p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # =========================================================================
    # MODELO DE NEGÓCIO
    # =========================================================================
    add_formatted_heading(doc, "MODELO DE NEGÓCIO", 1)
    doc.add_paragraph()
    
    add_formatted_heading(doc, "Fontes de Receita", 2)
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    headers = ["Produto", "Faixa de Preço", "Status"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '2E7D32')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    data = [
        ("Assinatura Mensal", "R$ 497 - 997/mês", "[NÃO VALIDADO]"),
        ("Honorários de Êxito", "10-15% do valor defendido", "[BENCHMARK]"),
        ("Consultoria Premium", "R$ 2.000 - 5.000/caso", "[ESTIMATIVA]")
    ]
    
    for i, (prod, preco, status) in enumerate(data, 1):
        table.rows[i].cells[0].text = prod
        table.rows[i].cells[1].text = preco
        table.rows[i].cells[2].text = status
        set_cell_shading(table.rows[i].cells[2], 'FFF3CD')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # =========================================================================
    # PROJEÇÕES FINANCEIRAS
    # =========================================================================
    add_formatted_heading(doc, "PROJEÇÕES FINANCEIRAS", 1)
    doc.add_paragraph()
    
    add_warning_box(doc, "TODAS as projeções abaixo são ESTIMATIVAS especulativas")
    
    # Tabela de projeções
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    
    headers = ["Métrica", "Ano 1", "Ano 2", "Ano 3", "Ano 4", "Ano 5"]
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], '2E7D32')
        table.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    
    projections = [
        ("Clientes", "150", "600", "2.000", "5.000", "10.000"),
        ("Receita", "R$ 1,2M", "R$ 6M", "R$ 24M", "R$ 72M", "R$ 180M"),
        ("EBITDA", "-R$ 800k", "R$ 600k", "R$ 7,2M", "R$ 28M", "R$ 72M")
    ]
    
    for i, row_data in enumerate(projections, 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val
            if j > 0:
                set_cell_shading(table.rows[i].cells[j], 'FFF3CD')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("⚠️ RESSALVAS CRÍTICAS:")
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    warnings = [
        "TIR projetada de ~100%+ é MUITO otimista para early-stage",
        "Projeções assumem execução perfeita e condições favoráveis",
        "Investidor deve aplicar desconto significativo a estes números",
        "Recomendamos solicitar cenários pessimista/base/otimista"
    ]
    
    for w in warnings:
        doc.add_paragraph(w, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # =========================================================================
    # EQUIPE
    # =========================================================================
    add_formatted_heading(doc, "EQUIPE", 1)
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run("[A PREENCHER COM DADOS REAIS]")
    run.italic = True
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    doc.add_paragraph()
    
    fields = [
        ("Fundador(es):", "_" * 50),
        ("Formação:", "_" * 50),
        ("Experiência relevante:", "_" * 50),
        ("E-mail:", "_" * 50),
        ("Telefone:", "_" * 50)
    ]
    
    for label, placeholder in fields:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        p.add_run(placeholder)
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # =========================================================================
    # INVESTIMENTO SOLICITADO
    # =========================================================================
    add_formatted_heading(doc, "INVESTIMENTO SOLICITADO", 1)
    doc.add_paragraph()
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    items = [
        ("Rodada:", "[Pré-Seed / Seed - DEFINIR]"),
        ("Valor:", "R$ [A DEFINIR]"),
        ("Valuation:", "R$ [A DEFINIR]"),
        ("Uso dos Recursos:", "[DETALHAR]")
    ]
    
    for i, (label, value) in enumerate(items):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], 'E8F5E9')
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    doc.add_page_break()
    
    # =========================================================================
    # PRÓXIMOS PASSOS
    # =========================================================================
    add_formatted_heading(doc, "PRÓXIMOS PASSOS PARA VALIDAÇÃO", 1)
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.add_run("Compromisso de transparência: ").bold = True
    p.add_run("Estamos comprometidos em validar todas as estimativas deste documento. Abaixo estão as ações planejadas:")
    
    doc.add_paragraph()
    
    steps = [
        "☐ Enviar pedidos LAI para IBAMA, SEMAs, MTE, MAPA",
        "☐ Entrevistar 20+ produtores para validar disposição a pagar",
        "☐ Confirmar dados de mercado com associações do setor (CNA, Aprosoja)",
        "☐ Desenvolver MVP e testar com 5-10 clientes piloto",
        "☐ Ajustar projeções com base em dados reais"
    ]
    
    for step in steps:
        p = doc.add_paragraph(step)
        p.paragraph_format.left_indent = Inches(0.3)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Rodapé final
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("─" * 60)
    run.font.color.rgb = RGBColor(200, 200, 200)
    
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final.add_run("Documento preparado com transparência para construir relação de confiança.\n")
    run.italic = True
    run = final.add_run("Versão: Janeiro 2026")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)
    
    return doc

if __name__ == "__main__":
    print("Gerando documento Word formatado...")
    doc = create_document()
    
    output_file = "Agrodefesa_Legal_Dossie_Transparente.docx"
    doc.save(output_file)
    
    print(f"✅ Documento gerado com sucesso: {output_file}")
    print(f"📁 Localização: C:\\Users\\Marcus Carvalho PC\\Documents\\AgroDefesa_Dossie\\{output_file}")
